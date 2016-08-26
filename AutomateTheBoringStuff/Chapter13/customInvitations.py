'''
Say you have a text file of guest names. This guests.txt file has one
name per line, as follows:

Prof. Plum
Miss Scarlet
Col. Mustard
Al Sweigart
Robocop

Write a program that would generate a Word document with custom invitations
that look like Figure 13-11.

Since Python-Docx can use only those styles that already exist in the Word
document, you will have to first add these styles to a blank Word file and
then open that file with Python-Docx. There should be one invitation per
page in the resulting Word document, so call add_break() to add a page
break after the last paragraph of each invitation. This way, you will need
to open only one Word document to print all of the invitations at once.
'''
import docx

lstAttendees = ['Prof. Plum','Miss Scarlet','Col. Mustard','Al Sweigart','Robocop']
lstWords = ['It would be a pleasure to have the company of','at 11010 Memory Lane on the Evening of','April 1st','at 7o\' clock']

doc = docx.Document('invitations.docx')

idx=0

for i in range(len(lstAttendees)):
    for v in range(len(lstWords)):
        if v==1:
            doc.add_paragraph(lstAttendees[i])
            idx+=1
            doc.paragraphs[idx].style='VeryNice2'
            doc.paragraphs[idx].runs[0].bold=True
        doc.add_paragraph(lstWords[v])
        idx+=1
        doc.paragraphs[idx].style='VeryNice'
        if v == (len(lstWords)-1):
            if i != len(lstAttendees):
                doc.paragraphs[idx].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

doc.save('invitations.docx')
